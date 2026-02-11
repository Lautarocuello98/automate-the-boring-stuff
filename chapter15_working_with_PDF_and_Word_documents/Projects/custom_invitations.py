# custom_invitations.py
# Generate a Word document with one invitation per guest name.

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

with open('guests.txt') as f:
    names = f.read().splitlines()

# Create a new Word document
doc = docx.Document()

# Loop through each guest name
for name in names:
    # Add the first line of the invitation using an existing style
    p = doc.add_paragraph("It would be a pleasure to have the company of", style="Quote")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Add the guest's name as a heading (large text)
    p = doc.add_heading(name, 1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("at 11010 Memory Lane on the Evening of", style="Quote")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Add the date using another existing style
    p = doc.add_paragraph("April 1st", style="Heading 2")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("at 7 o'clock", style="Quote")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Insert a page break
    doc.add_page_break()

doc.save("invitations.docx")