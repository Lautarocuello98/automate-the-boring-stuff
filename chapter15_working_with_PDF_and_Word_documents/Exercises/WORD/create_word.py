import docx

# Create doc
doc = docx.Document()

# Add headers
doc.add_heading('Header 0', 0)
doc.add_heading('Header 1', 1)
doc.add_heading('Header 2', 2)
doc.add_heading('Header 3', 3)
doc.add_heading('Header 4', 4)

# Add paragraphs
p = doc.add_paragraph('This is on the first page!')

# Page break
p.runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)

# Add other paragraphs
doc.add_paragraph('This is on the second page!')
doc.add_paragraph('This too is on the second page!')

# Add a picture
doc.add_picture('zophie.png', width=docx.shared.Inches(1))

doc.save('bigexample.docx')
