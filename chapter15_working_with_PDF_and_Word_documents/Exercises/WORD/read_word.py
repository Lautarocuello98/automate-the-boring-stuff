import docx

# Open the Word document
doc = docx.Document('demo.docx')

# Print how many paragraphs the document contains
print(len(doc.paragraphs))

# Print the text of the first paragraph
print(doc.paragraphs[0].text)

# Print the list of runs in the second paragraph
# (runs are pieces of text where the style doesn't change)
print(doc.paragraphs[1].runs)

# Print the text of the first run in the second paragraph
print(doc.paragraphs[1].runs[0].text)
