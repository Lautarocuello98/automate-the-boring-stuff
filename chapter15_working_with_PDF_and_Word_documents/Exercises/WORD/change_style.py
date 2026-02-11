import docx   # Import the library used to work with Word (.docx) files

# Open an existing document named 'demo.docx'
doc = docx.Document('demo.docx')

# Print the text of the first paragraph to the console
print(doc.paragraphs[0].text)

# Change the style of the first paragraph to the built-in Word style "Title"
doc.paragraphs[0].style = 'Title'

# Replace the entire text of the second paragraph
doc.paragraphs[1].text = ''

# Apply formatting to parts of the text in the second paragraph:
# runs[0] = first portion of text → set to bold
doc.paragraphs[1].add_run('A plain paragraph with some ')
doc.paragraphs[1].add_run('italic').italic = True

# runs[1] = second portion of text → set to italic
doc.paragraphs[1].add_run(' and some ')
doc.paragraphs[1].add_run('bold').bold = True

# Save the changes to a new file named 'restyled.docx'
doc.save('restyled.docx')
